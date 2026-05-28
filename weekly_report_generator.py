#!/usr/bin/env python3
"""Generate a formatted weekly report (.docx) from JSON data."""

import json
import platform
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is not installed. Run: pip3 install python-docx")
    sys.exit(1)

FONT_NAME = "SimSun"
FONT_FALLBACK = "STSong"


def get_font_name():
    """Return available CJK font with macOS fallback."""
    if platform.system() == "Darwin":
        return FONT_FALLBACK
    return FONT_NAME


def set_run_font(run, size_pt, bold=False):
    """Apply font settings to a run."""
    font = get_font_name()
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)


def add_title(doc, text):
    """Add centered bold title (二号 = 22pt)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 22, bold=True)


def add_centered_text(doc, text):
    """Add centered text (小四 = 12pt)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 12)


def add_heading_text(doc, text):
    """Add section heading (四号 = 14pt, bold)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)


def add_body_text(doc, text, indent=False):
    """Add body text (小四 = 12pt), optionally indented."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_paper_entry(doc, idx, paper):
    """Add a single paper entry."""
    title = paper.get("title", "Unknown Title")
    authors = paper.get("authors", "Unknown")
    year = paper.get("year", "N/A")
    citations = paper.get("citations", 0)
    summary = paper.get("research_summary", "")
    notes = paper.get("reading_notes", "")

    p = doc.add_paragraph()
    run = p.add_run(f"论文 {idx}：{title}")
    set_run_font(run, 12, bold=True)

    add_body_text(doc, f"作者：{authors} | 年份：{year} | 引用数：{citations}")
    add_body_text(doc, f"研究总结：{summary}", indent=True)
    add_body_text(doc, f"阅读笔记：{notes}", indent=True)
    doc.add_paragraph()


def generate_report(data, output_path):
    """Generate the .docx report from JSON data."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = get_font_name()
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    add_title(doc, data.get("title", "研究生周报"))

    date_range = data.get("date_range", {})
    start = date_range.get("start", "")
    end = date_range.get("end", "")
    add_centered_text(doc, f"时间：{start} - {end}")

    add_centered_text(doc, f"姓名：{data.get('name', '')}")

    doc.add_paragraph()

    add_heading_text(doc, "一、本周工作总结")
    add_body_text(doc, data.get("summary", ""), indent=True)

    add_heading_text(doc, "二、文献阅读")
    papers = data.get("papers", [])
    for idx, paper in enumerate(papers, 1):
        add_paper_entry(doc, idx, paper)

    add_heading_text(doc, "三、项目进展")
    for item in data.get("project_progress", []):
        add_body_text(doc, f"• {item}")

    add_heading_text(doc, "四、文献阅读总结")
    add_body_text(doc, data.get("paper_summary", ""), indent=True)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 weekly_report_generator.py <json_path> [output_path]")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "研究生周报.docx"

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: File not found: {json_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in {json_path}: {e}")
        sys.exit(1)

    output_dir = Path(output_path).parent
    if str(output_dir) != "." and not output_dir.exists():
        print(f"Error: Output directory does not exist: {output_dir}")
        sys.exit(1)

    result = generate_report(data, output_path)
    print(f"Report generated: {result}")


if __name__ == "__main__":
    main()
